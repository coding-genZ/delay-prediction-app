"""
Build the final project report — compact, professional, every page filled.
Generates FinalReport.docx then converts to PDF.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = os.path.join(os.path.dirname(__file__), "..")
FIG_DIR = os.path.join(ROOT, "figures")

doc = Document()
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x44, 0x44, 0x44)
WHITE_C = RGBColor(0xFF, 0xFF, 0xFF)

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def heading1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.font.color.rgb = BLACK
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)
    run2 = p2.add_run("_" * 95)
    run2.font.size = Pt(6)
    run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    return h


def heading2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = "Times New Roman"
    run.font.color.rgb = BLACK
    return h


def para(text, indent=True, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = space_after
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def bullet(text, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(f"•  {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE_C
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "333333")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            if ri % 2 == 0:
                set_cell_shading(cell, "F5F5F5")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return table


def add_image(path, width=Inches(5.5), caption=None):
    path = os.path.join(FIG_DIR, path) if not os.path.isabs(path) else path
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            r.font.color.rgb = GREY


# =====================================================================
# TITLE PAGE
# =====================================================================
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Shipment Delay Prediction App")
run.bold = True
run.font.size = Pt(26)
run.font.name = "Times New Roman"
run.font.color.rgb = BLACK

doc.add_paragraph().paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("A Proactive Risk Flagging System for E-Commerce Dispatch Operations")
run.font.size = Pt(13)
run.font.name = "Times New Roman"
run.font.color.rgb = GREY

for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = div.add_run("_" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.size = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(0)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Machine Learning for Business Applications\nFinal Project Report")
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run.font.color.rgb = BLACK

doc.add_paragraph().paragraph_format.space_after = Pt(0)

names = doc.add_paragraph()
names.alignment = WD_ALIGN_PARAGRAPH.CENTER
names.paragraph_format.space_before = Pt(12)
run = names.add_run(
    "Swarnadwip Bhattacharya\nHamna Asif\nNeha Pandey\nKerri Wang"
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run.font.color.rgb = BLACK

doc.add_paragraph().paragraph_format.space_after = Pt(0)

stack = doc.add_paragraph()
stack.alignment = WD_ALIGN_PARAGRAPH.CENTER
stack.paragraph_format.space_before = Pt(20)
run = stack.add_run("CatBoost  |  FastAPI  |  Streamlit  |  AWS Lambda  |  DynamoDB  |  S3  |  CloudWatch  |  SHAP")
run.font.size = Pt(9.5)
run.font.name = "Times New Roman"
run.font.color.rgb = GREY

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
heading1("Table of Contents")

toc = [
    ("1.", "Business Problem and Motivation", "3"),
    ("2.", "Dataset", "4"),
    ("3.", "Methodology", "7"),
    ("4.", "Results and Evaluation", "9"),
    ("5.", "Cloud Architecture", "12"),
    ("6.", "User Interface", "14"),
    ("7.", "Reflections and Lessons Learned", "16"),
    ("8.", "References", "18"),
]
for num, title, pg in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    r1 = p.add_run(f"{num}  {title}")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLACK

doc.add_page_break()

# =====================================================================
# 1. BUSINESS PROBLEM
# =====================================================================
heading1("1. Business Problem and Motivation")

para(
    "Late deliveries are among the most costly operational failures in e-commerce logistics. When a package "
    "misses its promised delivery window, the company faces a cascade of consequences: refunds, service recovery "
    "calls, discount coupons for the next order, and — worst of all — the long-term risk of losing a customer "
    "entirely. Industry estimates place the total cost of a single missed delivery promise at approximately $20 "
    "when accounting for the direct refund, customer service overhead, and projected lifetime value erosion."
)

para(
    "What makes this problem particularly frustrating is that most delays are predictable. The patterns are "
    "visible in the data: orders dispatched during peak traffic hours, assigned to couriers with historically low "
    "ratings, covering long distances in adverse weather — these shipments fail at dramatically higher rates. Yet "
    "most dispatch teams today rely on static rules or gut instinct when assigning couriers to orders. A common "
    'heuristic might be "avoid long-distance orders after 5 PM," but this misses the multivariate nature of '
    "delay risk and leaves money on the table."
)

para(
    "Our goal was to build a system that surfaces delay risk at the exact moment a dispatcher assigns a courier "
    "to an order. If the model flags a shipment as high-risk, the dispatcher can intervene — reassign to a "
    "better-rated courier, switch to a faster vehicle, adjust the route, or proactively notify the customer so "
    "they are not caught off guard. The end user is an operations team at an e-commerce fulfilment center that "
    "processes hundreds of courier assignments per hour and needs a fast, one-glance risk score per order."
)

para(
    "Critically, we wanted this to be more than a notebook exercise. The deliverable is a complete working "
    "system: a FastAPI backend serving real-time predictions with per-order SHAP explanations, a Streamlit "
    "frontend designed for dispatcher workflows, a full AWS cloud deployment using Lambda and API Gateway, and "
    "a feedback loop that captures dispatcher actions and actual outcomes to feed the next model version. Success "
    "means a dispatcher opens a URL, enters order details, and within two seconds sees a delay probability, a "
    "risk flag, the top three factors driving the score, and a concrete suggested action."
)

para(
    "The business value is measured in cost reduction. Every delay caught early and mitigated is $20 saved. "
    "Every false alarm costs $8 in unnecessary monitoring overhead. The model's job is to find the optimal "
    "trade-off between these two costs, and that trade-off is encoded directly into the classification threshold "
    "rather than left as an afterthought."
)

doc.add_page_break()

# =====================================================================
# 2. DATASET
# =====================================================================
heading1("2. Dataset")

heading2("2.1 Source and Size")

para(
    "We used the Amazon Delivery Dataset from Kaggle, a publicly available collection of 43,739 delivery records "
    "from an Indian e-commerce platform. Each row represents a single order and includes courier attributes, "
    "geographic coordinates for the store and drop-off, timestamps, weather and traffic conditions at the time "
    "of dispatch, and the actual delivery duration in minutes. After cleaning, 43,594 rows remain."
)

heading2("2.2 Features and Target Variable")

para(
    "After feature engineering, the model uses 12 input features across four categories:", indent=False
)

add_table(
    ["Category", "Features", "Type"],
    [
        ["Courier", "Agent age, agent rating", "Numeric"],
        ["Environment", "Weather (6 levels), traffic (4 levels)", "Categorical"],
        ["Logistics", "Vehicle type, area type, product category, distance (km)", "Mixed"],
        ["Temporal", "Pickup hour, day of week, weekend flag", "Numeric"],
    ],
    col_widths=[1.5, 3.5, 1.2]
)

para(
    "Distance was computed from store and drop-off coordinates using the haversine formula — it does not exist "
    "in the raw data. The weekend flag and pickup hour were extracted from order timestamps. For the target "
    'variable, we defined "delayed" as any delivery exceeding the median delivery time across the full dataset. '
    "This produced a near-balanced split of 47% delayed versus 53% on-time, eliminating the need for SMOTE or "
    "other resampling techniques."
)

add_image("class_distribution.png", width=Inches(4.0),
          caption="Figure 1: Target variable distribution — near-balanced classes (53% on-time, 47% delayed).")

heading2("2.3 Cleaning and Preprocessing")

para(
    "The raw dataset required several preprocessing steps before model training. We dropped rows "
    "with missing values in critical fields — primarily null delivery durations and unparseable timestamps. "
    "No imputation was applied because the missing rows appeared to be data collection failures (entire rows "
    "blank) rather than informative missingness. For the remaining rows, missingness was zero across "
    "all feature columns."
)

para(
    "Categorical features were left as raw strings for CatBoost, which handles encoding internally using "
    "ordered target statistics. Numeric features were not scaled or normalized — tree-based models are invariant "
    "to monotonic feature transformations, so standardization would add complexity without benefit. The haversine "
    "distance feature was computed from four coordinate columns, after which the raw latitude and longitude "
    "columns were dropped to avoid data leakage from coordinates into the model."
)

heading2("2.4 Data Quality")

para(
    "A key concern with Kaggle datasets is whether the data reflects real-world patterns or is synthetically "
    "generated with random labels. We validated the data in two ways. First, we checked that physical "
    "relationships hold: delivery time correlates with distance, high traffic increases delay rate, "
    "and stormy weather increases delay rate — all in the expected direction. Feature distributions also roughly "
    "match published Indian last-mile delivery benchmarks."
)

para(
    "However, we did observe that certain categorical features — particularly weather and product category — "
    "exhibit suspiciously uniform distributions, suggesting partial synthetic labeling. We carried this caveat "
    "throughout the project: the model’s primary value is in rank-ordering risk for dispatcher prioritization, "
    "not in treating the absolute probability as perfectly calibrated."
)

heading2("2.5 Exploratory Analysis")

para(
    "Before modeling, we examined how delay rate varies across the two most important features — traffic "
    "and weather — to confirm that the signal is strong enough for a classifier to learn meaningful patterns.",
    indent=False
)

add_image("delay_by_traffic.png", width=Inches(4.8),
          caption="Figure 2: Delay rate by traffic level — monotonic increase from Low (23.7%) to Jam (65.8%).")

para(
    "Traffic shows a clear monotonic relationship with delay rate. Orders dispatched in Jam conditions are "
    "nearly three times more likely to arrive late than those in Low traffic. This confirms Traffic as a "
    "first-order predictor and explains why it ranks highest in the global SHAP analysis."
)

add_image("delay_by_weather.png", width=Inches(4.8),
          caption="Figure 3: Delay rate by weather — Sunny is safest (23.9%), Cloudy and Fog above 55%.")

para(
    "Weather shows a clear difference between Sunny conditions (23.9% delay rate) and adverse conditions. "
    "Interestingly, Cloudy and Fog have higher delay rates than Stormy and Sandstorms, which suggests the "
    "near-uniform category distributions noted in Section 2.4 may be introducing noise. Despite this, weather "
    "still ranks fourth in global SHAP importance, providing meaningful signal."
)

heading2("2.6 Train/Test Split")

para(
    "We used an 80/20 stratified split. Stratification preserves "
    "the delayed/on-time ratio in both sets, which is important because our cost-sensitive threshold tuning "
    "(Section 3.3) requires the test set to be representative of the true class distribution."
)

doc.add_page_break()

# =====================================================================
# 3. METHODOLOGY
# =====================================================================
heading1("3. Methodology")

heading2("3.1 Model Selection")

para(
    "We benchmarked three classifiers on the same 80/20 split with identical features: logistic regression "
    "as a linear baseline, XGBoost as a strong gradient-boosted tree, and CatBoost as our primary candidate."
)

add_image("model_comparison.png", width=Inches(5.0),
          caption="Figure 4: Model comparison — CatBoost matches XGBoost on AUC and F1, leads on recall.")

add_table(
    ["Model", "ROC-AUC", "F1 (Delayed)", "Recall (Delayed)"],
    [
        ["Logistic Regression", "0.76", "0.67", "69%"],
        ["XGBoost", "0.95", "0.78", "86%"],
        ["CatBoost (Final)", "0.953", "0.781", "89.5%"],
    ],
    col_widths=[2.2, 1.3, 1.5, 1.5]
)

para(
    "CatBoost and XGBoost are nearly identical on ROC-AUC and F1, but CatBoost edges ahead on recall — catching "
    "89.5% of actual delays versus XGBoost's 86%. We selected CatBoost for two reasons. First, recall matters "
    "disproportionately for our cost structure: missing a delay costs $20 while a false alarm costs only $8, so "
    "a model that catches more delays at the cost of slightly more false alarms is more valuable. Second, "
    "CatBoost handles categorical features natively without one-hot encoding, which simplifies the inference "
    "pipeline and eliminates an entire class of production bugs around unseen categories."
)

heading2("3.2 Cross-Validation")

para(
    "A single train/test split can produce misleading results if the random split happens to be favorable. "
    "To verify stability, we ran 5-fold stratified cross-validation on the full dataset:"
)

bullet("Mean CV ROC-AUC: 0.9517 (standard deviation: 0.0025)")
bullet("Held-out test ROC-AUC: 0.9526")
bullet("Gap: 0.0009 — well within fold-to-fold variance")

para(
    "This confirms that the model is stable and not overfitting to a particular split. The held-out test score "
    "is consistent with what cross-validation predicts, not a lucky draw."
)

heading2("3.3 Cost-Sensitive Threshold Tuning")

para(
    "By default, classifiers use 0.5 as the decision boundary. This is rarely optimal for real business "
    "problems. In our case, the asymmetric cost structure — $20 per missed delay versus $8 per false alarm — "
    "means we should lower the threshold to catch more delays, accepting additional false alarms as a "
    "trade-off."
)

para(
    "We swept the threshold from 0.05 to 0.95 on the held-out test set, computing total business cost at "
    "each point: (missed delays × $20) + (false alarms × $8). The minimum cost occurs at threshold = 0.26."
)

add_image("threshold_cost_curve.png", width=Inches(5.0),
          caption="Figure 5: Business cost by classification threshold. Minimum at 0.26.")

para(
    "At this threshold, total cost on the test orders drops from approximately $13,000 (at default 0.5) to "
    "approximately $10,576 — an 18.6% reduction. This was the single most impactful improvement in the project. "
    "The model did not change; we simply aligned the decision boundary with the business cost structure."
)

heading2("3.4 Explainability with SHAP")

para(
    "A risk score without explanation is not actionable. If the model flags a 65% delay probability, the "
    "dispatcher needs to know why in order to decide what to do about it. We use SHAP (SHapley Additive "
    "exPlanations) to decompose every prediction into per-feature contributions."
)

para(
    "For each prediction, the API returns the top 3 features driving the risk score, their current values, "
    "whether they increase or decrease delay risk, and the magnitude (SHAP value). For example: \"Agent rating = "
    "2.6 (+4.36 SHAP), Traffic = Low (-0.99), Vehicle = motorcycle (+0.49).\" This tells the dispatcher that "
    "the low courier rating is the dominant factor and they should consider reassigning."
)

heading2("3.5 Hyperparameter Tuning")

para(
    "We built an Optuna-based Bayesian hyperparameter optimization pipeline (tune_optuna.py) that searches "
    "over eight CatBoost parameters: iterations (200–2000), tree depth (4–10), learning rate (0.01–0.3), "
    "L2 regularization (1–10), bagging temperature (0–10), random strength (0–10), border count (32–255), "
    "and minimum data in leaf (1–50). Each trial evaluates a configuration using 5-fold stratified "
    "cross-validation, maximizing mean ROC-AUC."
)

para(
    "The default CatBoost parameters already achieve 0.9526 AUC on the held-out test set. The tuning pipeline "
    "is designed to run as a production step — 50 trials by default, extensible to 200 for marginal gains. "
    "We include it as a ready-to-run script rather than reporting tuned results because the improvement over "
    "defaults is expected to be small (0.5–1% recall) and we prioritized building the full end-to-end system "
    "over squeezing the last fraction of model performance."
)

doc.add_page_break()

# =====================================================================
# 4. RESULTS
# =====================================================================
heading1("4. Results and Evaluation")

heading2("4.1 Final Performance Summary")

add_table(
    ["Metric", "Value", "Why It Matters"],
    [
        ["ROC-AUC", "0.9526", "Measures rank-ordering accuracy across all thresholds"],
        ["F1 (Delayed)", "0.781", "Balance of precision and recall at operating point"],
        ["Recall", "89.5%", "Fraction of actual delays the model catches"],
        ["Precision", "56.7%", "Fraction of flagged shipments that are actually delayed"],
        ["Threshold", "0.26", "Cost-optimal decision boundary"],
        ["CV ROC-AUC", "0.9517 ± 0.0025", "Confirms stability across 5 folds"],
        ["Cost Reduction", "18.6%", "~$13,000 → ~$10,576 on test orders"],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

para(
    "ROC-AUC is our primary evaluation metric because the dispatcher's core need is reliable risk ranking — not "
    "a perfect binary label, but a score that correctly orders shipments from lowest to highest delay risk. An "
    "AUC of 0.953 means the model ranks a randomly chosen delayed order above a randomly chosen on-time order "
    "95.3% of the time. F1 and recall at the tuned threshold quantify performance at the specific operating "
    "point we selected for the business."
)

heading2("4.2 ROC and Precision-Recall Curves")

para(
    "The ROC curve visualizes model performance across all possible thresholds. The strong concavity toward the "
    "top-left corner confirms excellent discrimination between delayed and on-time orders. The Precision-Recall "
    "curve (Average Precision = 0.92) is equally important for imbalanced-aware evaluation, showing that the "
    "model maintains high precision even at high recall levels.", indent=False
)

add_image("roc_curve.png", width=Inches(4.5),
          caption="Figure 6: ROC curve — CatBoost classifier (AUC = 0.91 on this split).")

add_image("precision_recall_curve.png", width=Inches(4.5),
          caption="Figure 7: Precision-Recall curve (Average Precision = 0.92).")

heading2("4.3 Confusion Matrix")

para(
    "The confusion matrix at the cost-optimal threshold of 0.26 on the held-out test set:", indent=False
)

add_image("confusion_matrix.png", width=Inches(4.5),
          caption="Figure 8: Confusion matrix at threshold 0.26 — 2,658 true positives, 165 false positives.")

para(
    "The model correctly identifies the large majority of actual delays as true positives. The false positives "
    "(on-time orders flagged as delayed) represent a manageable cost: a dispatcher who is alerted to monitor a "
    "shipment that turns out fine has wasted a few seconds of attention, not $20 in refunds. The false negatives "
    "(missed delays) are the costliest errors at $20 each, which is why our threshold is tuned to minimize them."
)

heading2("4.4 Score Distribution")

para(
    "The predicted probability distribution separated by true class shows how well the model separates the "
    "two populations. On-time orders cluster heavily near zero probability, while delayed orders spread across "
    "higher probabilities. The threshold line at 0.26 cleanly partitions the bulk of the on-time mass from "
    "the delayed mass.", indent=False
)

add_image("score_distribution.png", width=Inches(4.8),
          caption="Figure 9: Predicted probability distribution by true class, with threshold at 0.26.")

heading2("4.5 Global Feature Importance")

para(
    "The global SHAP analysis reveals which features drive delay predictions across the full dataset. "
    "The ranking by mean absolute SHAP value:", indent=False
)

add_table(
    ["Rank", "Feature", "Mean |SHAP|"],
    [
        ["1", "Traffic", "1.32"],
        ["2", "Agent Age", "1.22"],
        ["3", "Agent Rating", "0.84"],
        ["4", "Weather", "0.71"],
        ["5", "Vehicle", "0.60"],
        ["6", "Distance (km)", "0.57"],
        ["7", "Category", "0.52"],
        ["8", "Area", "0.12"],
        ["9", "Pickup Hour", "0.10"],
        ["10", "Day of Week", "0.03"],
        ["11", "Prep Minutes", "0.02"],
        ["12", "Is Weekend", "0.01"],
    ],
    col_widths=[0.7, 2.0, 1.5]
)

add_image("shap_importance_bar.png", width=Inches(5.2),
          caption="Figure 10: Global feature importance by mean absolute SHAP value.")

para(
    "Traffic is the single strongest predictor, followed by courier age and rating. This aligns with last-mile "
    "delivery research — road conditions and courier quality are the dominant factors. Notably, weather ranks "
    "fourth, lower than most dispatchers would intuit. The temporal features (pickup hour, day of week, weekend) "
    "have very low importance because traffic level already captures most of that signal — they are correlated, "
    "and the model gives credit to the more directly predictive feature."
)

add_image("shap_summary.png", width=Inches(5.2),
          caption="Figure 11: SHAP summary plot showing feature value impact direction and magnitude.")

para(
    "The SHAP summary plot above shows how individual feature values push predictions. Red dots represent high "
    "feature values, blue dots represent low values. For Agent Age, high values (older couriers) push predictions "
    "toward delay, while for Agent Rating, low values (poorly rated couriers) increase delay risk — visible as "
    "the blue cluster on the right side of the Agent Rating row."
)

doc.add_page_break()

# =====================================================================
# 5. CLOUD ARCHITECTURE
# =====================================================================
heading1("5. Cloud Architecture")

heading2("5.1 System Overview")

para(
    "The application runs on six AWS services, all provisioned through a single SAM (Serverless Application "
    "Model) template. The architecture is designed for serverless, pay-per-use operation with zero idle cost.",
    indent=False
)

# Architecture diagram as table
arch_diagram = doc.add_table(rows=5, cols=5)
arch_diagram.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_diagram.style = "Table Grid"

for cell in arch_diagram.rows[0].cells:
    set_cell_shading(cell, "333333")

arch_labels = [
    (0, 0, "Streamlit UI\n(User Input)"),
    (0, 1, "→"),
    (0, 2, "API Gateway\n(HTTPS)"),
    (0, 3, "→"),
    (0, 4, "Lambda\n(FastAPI + CatBoost)"),
    (1, 0, ""),
    (1, 1, ""),
    (1, 2, ""),
    (1, 3, ""),
    (1, 4, "↕"),
    (2, 0, "OpenWeatherMap\n(Live Weather)"),
    (2, 1, "→"),
    (2, 2, "Lambda"),
    (2, 3, "→"),
    (2, 4, "S3\n(Model Storage)"),
    (3, 0, ""),
    (3, 1, ""),
    (3, 2, "↓"),
    (3, 3, ""),
    (3, 4, ""),
    (4, 0, "Retrain Pipeline"),
    (4, 1, "←"),
    (4, 2, "DynamoDB\n(Predictions +\nFeedback)"),
    (4, 3, "←"),
    (4, 4, "CloudWatch\n(Logging)"),
]
for r, c, text in arch_labels:
    cell = arch_diagram.rows[r].cells[c]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.color.rgb = BLACK if r > 0 else WHITE_C
    if r == 0:
        run.bold = True

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_after = Pt(8)
r = cap.add_run("Figure 12: System architecture — prediction and feedback data flow.")
r.italic = True
r.font.size = Pt(9)
r.font.name = "Times New Roman"
r.font.color.rgb = GREY

add_table(
    ["AWS Service", "Role in the System"],
    [
        ["S3", "Stores trained model (.cbm) and config (model_info.json). New models uploaded after retraining."],
        ["Lambda", "Runs FastAPI via Mangum adapter. Loads model from S3 on cold start, serves predictions, computes SHAP."],
        ["API Gateway", "HTTP API providing a public HTTPS endpoint. Routes requests to Lambda with CORS support."],
        ["DynamoDB", "Stores every prediction (request + response + SHAP) and dispatcher feedback (action + outcome)."],
        ["CloudWatch", "Centralized structured logging for all predictions, errors, and feedback events."],
        ["Streamlit", "Dispatcher-facing web UI. Runs locally or on Streamlit Community Cloud."],
    ],
    col_widths=[1.5, 5.0]
)

heading2("5.2 Prediction Flow")

para("When a dispatcher clicks \"Predict Delay Risk,\" the following sequence executes:", indent=False)

bullet("Streamlit sends a POST request with order details to the API Gateway endpoint.")
bullet("API Gateway routes the request to the Lambda function.")
bullet("Lambda loads the CatBoost model from S3 (cold start) or memory (warm start).")
bullet("If configured, the system fetches live weather from OpenWeatherMap at both pickup and drop-off coordinates, overriding the manually entered weather value with real-time data.")
bullet("CatBoost predicts the delay probability. SHAP TreeExplainer computes per-feature contributions.")
bullet("The full prediction (request, response, SHAP values, weather data) is logged to DynamoDB with a unique prediction ID.")
bullet("The response — probability, risk flag, top 3 drivers, suggested action — returns to Streamlit in under 2 seconds.")

heading2("5.3 Feedback Loop")

para(
    "The architectural centerpiece is a closed feedback loop. Every prediction receives a unique ID. After "
    "delivery, the dispatcher records what action they took (reassigned courier, notified customer, no action, "
    "etc.) and the actual outcome (on-time, late, cancelled). This tuple — (order, prediction, action, "
    "outcome) — is stored in DynamoDB alongside the original prediction."
)

para(
    "A retraining pipeline (sagemaker_retrain.py) pulls all completed feedback records from DynamoDB, merges "
    "them with the original training data, retrains CatBoost, and — if the new model's AUC exceeds a "
    "configurable guard rail — uploads it to S3. Lambda picks up the new model on the next cold start. This "
    "loop is designed to run monthly via an EventBridge scheduled rule. Over time, this feedback data becomes "
    "the most valuable asset in the system because the model learns from real dispatcher decisions and actual "
    "delivery outcomes, not just historical patterns."
)

heading2("5.4 Deployment")

para(
    "The entire infrastructure is defined in template.yaml. Deployment requires two commands:", indent=False
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("sam build && sam deploy --guided")
run.bold = True
run.font.name = "Courier New"
run.font.size = Pt(10)

para(
    "SAM creates the Lambda function, API Gateway, DynamoDB table, S3 bucket, and CloudWatch log group. The "
    "output prints the public API URL. For local development, every AWS integration is gated behind environment "
    "variables and degrades gracefully — if S3_MODEL_BUCKET is not set, the model loads from local files; if "
    "DynamoDB is not configured, predictions are not logged; if CloudWatch is not set, logs go to stdout. A "
    "developer can run the full application on their laptop with zero AWS configuration."
)

doc.add_page_break()

# =====================================================================
# 6. USER INTERFACE
# =====================================================================
heading1("6. User Interface")

para(
    "The frontend is a Streamlit application optimized for dispatcher workflows. The design prioritizes speed "
    "and clarity — a dispatcher processing hundreds of orders per hour needs to glance at the screen and "
    "immediately know whether to intervene on a given shipment."
)

heading2("6.1 Order Scoring View")

para(
    "The main screen presents three input columns: Courier Details (age, rating, vehicle type), Route "
    "Information (store and drop-off coordinates, area type), and Conditions & Timing (weather, traffic, "
    "pickup hour, day of week, prep time, product category). After entering order details, the dispatcher "
    "clicks a single button to get the prediction."
)

add_image("ui_prediction.png", width=Inches(5.5),
          caption="Figure 13: Order input form — courier details, route information, and conditions.")

para(
    "Results appear as three color-coded metric cards — delay probability (blue), risk flag (red for DELAYED, "
    "green for ON-TIME), and the decision threshold (yellow). Below these cards, a gradient risk bar provides "
    'a visual anchor: green for low risk, yellow for elevated, red for high. A recommendation box suggests '
    'a specific action based on the risk tier: "Low risk — no action needed," "Elevated risk — monitor '
    'shipment; consider route/ETA adjustment," or "High risk — consider reassigning courier or notifying '
    'customer proactively."'
)

para(
    "The SHAP drivers table shows the top 3 features behind the prediction, their values, impact direction "
    "(increases/decreases risk as colored badges), and the SHAP magnitude. This transforms the model from a "
    "black-box number into an actionable explanation."
)

add_image("UI_Prediction result .png", width=Inches(5.5),
          caption="Figure 14: Prediction result — risk cards, risk bar, action recommendation, and SHAP explanation table.")

heading2("6.2 Dispatcher Feedback")

para(
    "Below every prediction result, a feedback form lets the dispatcher record two pieces of information: "
    "the action they took (no action, reassigned courier, notified customer, adjusted route, escalated, other) "
    "and the actual delivery outcome (delivered on time, delivered late, cancelled, returned). An optional notes "
    "field captures qualitative context. This data is submitted to the /feedback API endpoint and stored in "
    "DynamoDB, creating the training set for the next model version."
)

add_image("UI_feedback.png", width=Inches(5.5),
          caption="Figure 15: Full prediction view with dispatcher feedback form — action taken, outcome, and notes.")

heading2("6.3 Live Weather Integration")

para(
    "When an OpenWeatherMap API key is configured, the system automatically fetches current weather at both "
    "pickup and drop-off coordinates. A blue banner displays the live conditions — temperature, wind speed, "
    "visibility — and the weather category used by the model. This overrides the dispatcher's manual weather "
    "selection, giving the model a more accurate and objective signal. The integration is transparent: the "
    'banner clearly states that weather was "automatically fetched from OpenWeatherMap."'
)

heading2("6.4 Model Information Page")

para(
    'A sidebar toggle switches to an "About the Model" view displaying key metrics (ROC-AUC, F1, threshold), '
    "the cost-sensitive tuning rationale, the data quality caveat, and evaluation figures. This provides "
    "transparency for dispatchers or managers who want to understand the model's confidence level before "
    "relying on its recommendations."
)

doc.add_page_break()

# =====================================================================
# 7. REFLECTIONS
# =====================================================================
heading1("7. Reflections and Lessons Learned")

heading2("7.1 What Worked Well")

para(
    "CatBoost's native categorical feature handling was the right architectural choice. We spent virtually no "
    "time on feature encoding — no one-hot encoding, no label encoding, no worrying about unseen categories at "
    "inference time. The model takes raw strings and handles them internally. For a project where we needed to "
    "reach a working end-to-end application quickly, this eliminated an entire category of preprocessing work "
    "and production bugs."
)

para(
    "Cost-sensitive threshold tuning was the single most impactful technique we applied. The model's "
    "discriminative power did not change, but aligning the decision boundary with the actual business cost "
    "structure ($20 per missed delay vs. $8 per false alarm) reduced total cost by 18.6%. If we had to "
    "recommend one technique from this project to apply to a different classification problem, it would be "
    "threshold tuning — it is simple, requires no retraining, and the improvement is immediate."
)

para(
    "Cleanly separating training from serving paid dividends throughout the project. The model trains once and "
    "saves as a .cbm file. The FastAPI server loads it at startup and serves predictions. When we later added "
    "Lambda deployment, S3 model loading, DynamoDB logging, live weather enrichment, and a feedback endpoint, "
    "the core prediction code required minimal changes. The separation of concerns also made the SAM deployment "
    "straightforward — the same api.py file runs locally with Uvicorn and in Lambda with Mangum."
)

heading2("7.2 What Was Harder Than Expected")

para(
    "SHAP integration was the most time-consuming challenge. The SHAP library's TreeExplainer works well "
    "but initial calls took 20-30 seconds because of an initialization step we had not anticipated. The fix "
    "was straightforward — initialize the explainer once at startup rather than per-request — but diagnosing "
    "the bottleneck took longer than the actual fix."
)

para(
    "OneDrive syncing caused repeated Python venv corruption on Windows. The virtual environment's file locks "
    "conflicted with OneDrive's sync mechanism, breaking pip and package imports unpredictably. We had to be "
    "deliberate about directory placement and eventually added this to our lessons learned for future projects "
    "on Windows."
)

para(
    "Lambda cold starts with CatBoost were slower than expected. The model file is approximately 800 KB (small "
    "for a gradient-boosted model), but loading it plus initializing the SHAP explainer means the first request "
    "after a cold start takes several seconds. Allocating 1024 MB of memory to the Lambda function brought "
    "this to an acceptable level, but it remains a consideration for production latency SLAs."
)

heading2("7.3 What We Would Do Differently")

para(
    "We would implement the feedback loop on day one and begin collecting dispatcher feedback immediately, even "
    "before the model was fully tuned. In a real deployment, the feedback dataset — (order, prediction, "
    "dispatcher action, actual outcome) — is more valuable than the model itself, because models can be "
    "retrained but ground-truth operational labels are difficult and expensive to obtain retroactively."
)

para(
    "We would also invest more in the weather signal. The current model uses categorical weather labels (Sunny, "
    "Cloudy, Stormy). We built a live weather integration with OpenWeatherMap, but it still maps to these same "
    "categories. Adding continuous weather features — temperature, wind speed, visibility, precipitation rate — "
    "as direct model inputs would likely capture the residual error on weather-driven delays that the categorical "
    "labels miss."
)

para(
    "We would also explore temporal features more deeply. Delivery demand follows strong weekly and seasonal "
    "patterns — festival seasons, end-of-month spikes, weekend vs. weekday dynamics. Incorporating these "
    "cyclical patterns through sine/cosine encoding or as external calendar features could improve the model's "
    "ability to anticipate demand-driven congestion that is not captured by the current traffic variable."
)

heading2("7.4 Planned Next Steps")

bullet("Run Optuna Bayesian hyperparameter tuning. The script (tune_optuna.py) is built and ready. We expect a 0.5-1% recall improvement from systematic search over the current default parameters.")
bullet("Schedule monthly automated retraining via EventBridge. The retraining pipeline (sagemaker_retrain.py) is complete; it needs to be wired to a scheduled trigger.")
bullet("Execute the bias audit (bias_audit.py) before any real production deployment. The script evaluates whether model predictions exhibit disparate impact across courier age groups and rating bands — a fundamental fairness check.")
bullet("Activate live weather enrichment with an OpenWeatherMap API key and measure whether it improves accuracy on days with severe weather.")
bullet("Host Streamlit publicly on Streamlit Community Cloud (free tier) or a small EC2 instance behind CloudFront, so dispatchers access the app from a URL rather than running Python locally.")
bullet("Add A/B testing infrastructure to measure the causal impact of model-guided dispatch decisions on delivery times, comparing routes where dispatchers acted on predictions versus control routes.")

doc.add_page_break()

# =====================================================================
# 8. REFERENCES
# =====================================================================
heading1("8. References")

refs = [
    ("1", "Amazon Delivery Dataset. Kaggle. Available at: https://www.kaggle.com/datasets"),
    ("2", "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A.V., & Gulin, A. (2018). \"CatBoost: Unbiased Boosting with Categorical Features.\" Advances in Neural Information Processing Systems (NeurIPS), 31."),
    ("3", "Lundberg, S.M. & Lee, S.I. (2017). \"A Unified Approach to Interpreting Model Predictions.\" Advances in Neural Information Processing Systems (NeurIPS), 30."),
    ("4", "AWS Serverless Application Model (SAM) Developer Guide. Amazon Web Services. https://docs.aws.amazon.com/sam/"),
    ("5", "Streamlit Documentation. Streamlit Inc. https://docs.streamlit.io/"),
    ("6", "FastAPI Documentation. Tiangolo. https://fastapi.tiangolo.com/"),
    ("7", "OpenWeatherMap Current Weather Data API. OpenWeather Ltd. https://openweathermap.org/api"),
    ("8", "Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). \"Optuna: A Next-generation Hyperparameter Optimization Framework.\" Proceedings of ACM SIGKDD, 2623-2631."),
    ("9", "Chen, T. & Guestrin, C. (2016). \"XGBoost: A Scalable Tree Boosting System.\" Proceedings of ACM SIGKDD, 785-794."),
    ("10", "Barocas, S. & Selbst, A.D. (2016). \"Big Data's Disparate Impact.\" California Law Review, 104(3), 671-732."),
]

for num, ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    r = p.add_run(f"[{num}]  {ref}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK

# ── Save ──
output = os.path.join(ROOT, "FinalReport.docx")
doc.save(output)
print(f"Saved: {output}")
